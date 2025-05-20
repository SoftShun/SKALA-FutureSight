import os
from pathlib import Path
from typing import Optional
import tempfile
import datetime

# PDF 변환용
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

# DOCX 변환용
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 마크다운 파싱용
import markdown
from bs4 import BeautifulSoup

class ReportConverter:
    """
    보고서 변환 도구: 마크다운에서 PDF, DOCX 형식으로 변환
    """
    
    def __init__(self):
        """ReportConverter 초기화"""
        self.output_dir = Path("output")
        
        # 출력 디렉토리가 없으면 생성
        if not self.output_dir.exists():
            self.output_dir.mkdir(parents=True)
    
    def convert(
        self, 
        markdown_content: str, 
        format: str = "md", 
        filename: str = "report"
    ) -> str:
        """
        마크다운을 지정된 형식으로 변환
        
        Args:
            markdown_content: 마크다운 형식의 보고서 내용
            format: 출력 형식 (md, pdf, docx)
            filename: 출력 파일 이름 (확장자 제외)
            
        Returns:
            str: 변환된 파일의 경로
        """
        if format == "md":
            return self._save_markdown(markdown_content, filename)
        elif format == "pdf":
            return self._convert_to_pdf(markdown_content, filename)
        elif format == "docx":
            return self._convert_to_docx(markdown_content, filename)
        else:
            raise ValueError(f"지원되지 않는 형식: {format}")
    
    def _save_markdown(self, content: str, filename: str) -> str:
        """마크다운 파일로 저장"""
        output_path = self.output_dir / f"{filename}.md"
        
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(content)
        
        return str(output_path)
    
    def _convert_to_pdf(self, markdown_content: str, filename: str) -> str:
        """마크다운을 PDF로 변환"""
        output_path = self.output_dir / f"{filename}.pdf"
        
        # 한글 폰트 등록
        gothic_font_path = "/System/Library/Fonts/Supplemental/AppleGothic.ttf"
        pdfmetrics.registerFont(TTFont("AppleGothic", gothic_font_path))
        
        # 마크다운을 HTML로 변환
        html = markdown.markdown(markdown_content, extensions=['tables', 'fenced_code'])
        
        # PDF 생성
        doc = SimpleDocTemplate(str(output_path), pagesize=A4)
        styles = getSampleStyleSheet()
        flowables = []
        
        # 한글을 위한 스타일 정의
        korean_style = ParagraphStyle(
            'KoreanStyle',
            parent=styles['Normal'],
            fontName='AppleGothic',
            fontSize=10,
            leading=14,
            spaceAfter=10
        )
        
        # 한글을 위한 제목 스타일 정의
        korean_heading1 = ParagraphStyle(
            'KoreanHeading1',
            parent=styles['Heading1'],
            fontName='AppleGothic',
            fontSize=16,
            leading=20,
            spaceAfter=12
        )
        
        korean_heading2 = ParagraphStyle(
            'KoreanHeading2',
            parent=styles['Heading2'],
            fontName='AppleGothic',
            fontSize=14,
            leading=18,
            spaceAfter=10
        )
        
        # HTML 파싱하여 처리
        soup = BeautifulSoup(html, 'html.parser')
        
        # HTML 구조에서 텍스트 추출 및 적절한 스타일 적용
        for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'ul', 'ol', 'li']):
            text = element.get_text().strip()
            if not text:
                continue
                
            if element.name == 'h1':
                p = Paragraph(text, korean_heading1)
            elif element.name == 'h2':
                p = Paragraph(text, korean_heading2)
            elif element.name in ['h3', 'h4', 'h5', 'h6']:
                p = Paragraph(text, korean_heading2)
            else:
                p = Paragraph(text, korean_style)
                
            flowables.append(p)
            flowables.append(Spacer(1, 6))
        
        doc.build(flowables)
        return str(output_path)
    
    def _convert_to_docx(self, markdown_content: str, filename: str) -> str:
        """마크다운을 DOCX로 변환"""
        output_path = self.output_dir / f"{filename}.docx"
        
        # 새 Word 문서 생성
        doc = Document()
        
        # 기본 스타일 설정 - 한글 지원 범용 폰트로 변경
        style = doc.styles['Normal']
        style.font.name = 'Malgun Gothic'  # Windows의 경우 '맑은 고딕', Mac의 경우 AppleGothic과 호환
        style.font.size = Pt(11)
        
        # 마크다운 내용을 단순화하여 처리
        lines = markdown_content.split("\n")
        current_paragraph = None
        
        for line in lines:
            line = line.rstrip()
            
            # 제목 처리
            if line.startswith('# '):
                heading = doc.add_heading(line[2:], level=1)
                for run in heading.runs:
                    run.font.name = 'Malgun Gothic'  # 제목도 한글 폰트로 설정
            elif line.startswith('## '):
                heading = doc.add_heading(line[3:], level=2)
                for run in heading.runs:
                    run.font.name = 'Malgun Gothic'
            elif line.startswith('### '):
                heading = doc.add_heading(line[4:], level=3)
                for run in heading.runs:
                    run.font.name = 'Malgun Gothic'
            # 빈 줄 처리
            elif not line:
                current_paragraph = None
            # 일반 텍스트 처리
            else:
                if current_paragraph is None:
                    current_paragraph = doc.add_paragraph()
                run = current_paragraph.add_run(line)
                run.font.name = 'Malgun Gothic'  # 각 텍스트 블록마다 폰트 설정
        
        # 문서 저장
        doc.save(str(output_path))
        return str(output_path) 